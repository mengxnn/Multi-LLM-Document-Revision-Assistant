from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document


HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")


@dataclass(frozen=True)
class PdfTextExtraction:
    text: str
    page_layouts: tuple[str, ...]

    @property
    def used_two_column_layout(self) -> bool:
        return "two-column" in self.page_layouts


@dataclass(frozen=True)
class _PdfTextBlock:
    x0: float
    y0: float
    x1: float
    y1: float
    text: str

    @property
    def center_x(self) -> float:
        return (self.x0 + self.x1) / 2

    @property
    def center_y(self) -> float:
        return (self.y0 + self.y1) / 2

def read_source_text(path: str | Path) -> str:
    source_path = Path(path)
    suffix = source_path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return source_path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return read_docx_text(source_path)
    if suffix == ".pdf":
        return read_pdf_text(source_path)
    if suffix == ".doc":
        raise ValueError(".doc is not supported yet. Please convert it to .docx first.")
    raise ValueError(f"Unsupported source file type: {source_path.suffix}")


def read_docx_text(path: str | Path) -> str:
    document = Document(path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        heading_level = _heading_level(paragraph.style.name)
        if heading_level:
            blocks.append(f"{'#' * heading_level} {text}")
        else:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [_clean_cell_text(cell.text) for cell in row.cells]
            if any(cells):
                blocks.append("| " + " | ".join(cells) + " |")

    return "\n\n".join(blocks)


def read_pdf_text(path: str | Path) -> str:
    return extract_pdf_text(path).text


def extract_pdf_text(path: str | Path) -> PdfTextExtraction:
    document = fitz.open(str(path))
    pages: list[str] = []
    page_layouts: list[str] = []
    try:
        for index, page in enumerate(document, start=1):
            blocks = _pdf_text_blocks(page)
            if not blocks:
                page_layouts.append("empty")
                continue
            two_column = _is_two_column_page(
                blocks,
                page_width=float(page.rect.width),
                page_height=float(page.rect.height),
            )
            page_layouts.append("two-column" if two_column else "single-column")
            if two_column:
                ordered = _order_two_column_blocks(
                    blocks,
                    page_width=float(page.rect.width),
                )
            else:
                ordered = sorted(blocks, key=lambda item: (item.y0, item.x0))
            text = _join_pdf_blocks(ordered)
            if text:
                pages.append(f"<!-- page {index} -->\n\n{text}")
    finally:
        document.close()
    if not pages:
        raise ValueError(
            "No text could be extracted from this PDF. It may be scanned or image-only and needs OCR."
        )
    return PdfTextExtraction(
        text="\n\n".join(pages),
        page_layouts=tuple(page_layouts),
    )


def _pdf_text_blocks(page) -> list[_PdfTextBlock]:
    blocks: list[_PdfTextBlock] = []
    for raw_block in page.get_text("dict", sort=False).get("blocks", []):
        if int(raw_block.get("type", 0)) != 0:
            continue
        for line in raw_block.get("lines", []):
            text = "".join(
                str(span.get("text") or "")
                for span in line.get("spans", [])
            ).strip()
            bbox = line.get("bbox")
            if not text or not bbox or len(bbox) < 4:
                continue
            blocks.append(
                _PdfTextBlock(
                    x0=float(bbox[0]),
                    y0=float(bbox[1]),
                    x1=float(bbox[2]),
                    y1=float(bbox[3]),
                    text=text,
                )
            )
    return blocks


def _is_two_column_page(
    blocks: list[_PdfTextBlock],
    *,
    page_width: float,
    page_height: float,
) -> bool:
    split = page_width / 2
    center_tolerance = page_width * 0.04
    left = [
        block
        for block in blocks
        if block.center_x < split and block.x1 <= split + center_tolerance
    ]
    right = [
        block
        for block in blocks
        if block.center_x >= split and block.x0 >= split - center_tolerance
    ]
    if sum(len(block.text) for block in left) < 40:
        return False
    if sum(len(block.text) for block in right) < 40:
        return False

    left_center = _weighted_column_center(left)
    right_center = _weighted_column_center(right)
    if right_center - left_center < page_width * 0.25:
        return False

    overlap = min(
        max(block.y1 for block in left),
        max(block.y1 for block in right),
    ) - max(
        min(block.y0 for block in left),
        min(block.y0 for block in right),
    )
    return overlap >= max(12.0, page_height * 0.02)


def _weighted_column_center(blocks: list[_PdfTextBlock]) -> float:
    weights = [max(len(block.text), 1) for block in blocks]
    weighted_sum = sum(
        block.center_x * weight
        for block, weight in zip(blocks, weights)
    )
    return weighted_sum / sum(weights)


def _order_two_column_blocks(
    blocks: list[_PdfTextBlock],
    *,
    page_width: float,
) -> list[_PdfTextBlock]:
    split = page_width / 2
    separators = sorted(
        (
            block
            for block in blocks
            if block.x0 < split - page_width * 0.01
            and block.x1 > split + page_width * 0.01
        ),
        key=lambda item: (item.y0, item.x0),
    )
    pending = [block for block in blocks if block not in separators]
    ordered: list[_PdfTextBlock] = []
    for separator in separators:
        before = [block for block in pending if block.center_y < separator.center_y]
        ordered.extend(_order_column_band(before, split=split))
        ordered.append(separator)
        pending = [block for block in pending if block not in before]
    ordered.extend(_order_column_band(pending, split=split))
    return ordered


def _order_column_band(
    blocks: list[_PdfTextBlock],
    *,
    split: float,
) -> list[_PdfTextBlock]:
    left = sorted(
        (block for block in blocks if block.center_x < split),
        key=lambda item: (item.y0, item.x0),
    )
    right = sorted(
        (block for block in blocks if block.center_x >= split),
        key=lambda item: (item.y0, item.x0),
    )
    return [*left, *right]


def _join_pdf_blocks(blocks: list[_PdfTextBlock]) -> str:
    if not blocks:
        return ""
    parts = [blocks[0].text]
    for previous, current in zip(blocks, blocks[1:]):
        vertical_gap = current.y0 - previous.y1
        line_height = max(previous.y1 - previous.y0, current.y1 - current.y0)
        horizontal_overlap = min(previous.x1, current.x1) - max(
            previous.x0,
            current.x0,
        )
        same_text_flow = (
            current.y0 >= previous.y0
            and vertical_gap <= line_height * 0.75
            and horizontal_overlap > 0
        )
        parts.append(("\n" if same_text_flow else "\n\n") + current.text)
    return "".join(parts).strip()


def write_final_docx(text: str, output_path: str | Path, reference_path: str | Path | None = None) -> None:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document(reference_path) if reference_path else Document()
    if reference_path:
        _clear_body(document)

    pending_table: list[list[str]] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            _flush_table(document, pending_table)
            pending_table = []
            continue
        if _is_horizontal_rule(line):
            _flush_table(document, pending_table)
            pending_table = []
            continue
        if _is_table_row(line):
            pending_table.append(_parse_table_row(line))
            continue

        _flush_table(document, pending_table)
        pending_table = []
        heading = HEADING_PATTERN.match(line)
        if heading:
            level = min(len(heading.group(1)), 6)
            paragraph = document.add_heading("", level=level)
            _add_markdown_runs(paragraph, heading.group(2).strip())
        else:
            paragraph = document.add_paragraph()
            _add_markdown_runs(paragraph, line)

    _flush_table(document, pending_table)
    document.save(output)


def _heading_level(style_name: str) -> int | None:
    match = re.search(r"Heading\s+([1-6])", style_name, flags=re.IGNORECASE)
    if not match:
        return None
    return int(match.group(1))


def _clean_cell_text(text: str) -> str:
    return " ".join(text.split())


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def _is_horizontal_rule(line: str) -> bool:
    return bool(re.fullmatch(r"[-*_]{3,}", line.replace(" ", "")))


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip("|").split("|")]


def _flush_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    rows = [row for row in rows if not _is_separator_row(row)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_markdown_runs(paragraph, row[col_index] if col_index < len(row) else "")


def _is_separator_row(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row if cell.strip())


def _add_markdown_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|<br\s*/?>)", text, flags=re.IGNORECASE)
    for part in parts:
        if not part:
            continue
        if re.fullmatch(r"<br\s*/?>", part, flags=re.IGNORECASE):
            paragraph.add_run().add_break()
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part.replace("\\n", "\n"))


def _clear_body(document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)
