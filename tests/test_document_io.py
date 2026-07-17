import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from office_revision.document_io import (
    extract_pdf_text,
    read_source_text,
    write_final_docx,
)


def write_text_pdf(path: Path, text: str) -> None:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(output)


def write_blank_pdf(path: Path) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with path.open("wb") as output:
        writer.write(output)


def write_layout_pdf(path: Path, *, two_columns: bool) -> None:
    import fitz

    document = fitz.open()
    page = document.new_page(width=600, height=800)
    page.insert_textbox(
        fitz.Rect(50, 35, 550, 80),
        "FULL WIDTH PAPER TITLE",
        fontsize=14,
        align=fitz.TEXT_ALIGN_CENTER,
    )
    if two_columns:
        # Insert the right column first to ensure extraction is based on layout,
        # not PDF object order.
        page.insert_textbox(
            fitz.Rect(330, 120, 560, 300),
            "RIGHT COLUMN FIRST PARAGRAPH.\nRIGHT COLUMN SECOND PARAGRAPH.",
            fontsize=11,
        )
        page.insert_textbox(
            fitz.Rect(40, 120, 270, 300),
            "LEFT COLUMN FIRST PARAGRAPH.\nLEFT COLUMN SECOND PARAGRAPH.",
            fontsize=11,
        )
    else:
        page.insert_textbox(
            fitz.Rect(70, 120, 530, 300),
            "SINGLE COLUMN FIRST PARAGRAPH.\nSINGLE COLUMN SECOND PARAGRAPH.",
            fontsize=11,
        )
    page.insert_textbox(
        fitz.Rect(50, 700, 550, 740),
        "FULL WIDTH FOOTER",
        fontsize=10,
        align=fitz.TEXT_ALIGN_CENTER,
    )
    document.save(path)
    document.close()


class DocumentIoTests(unittest.TestCase):
    def test_reads_text_pdf_pages(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "source.pdf"
            write_text_pdf(path, "PDF requirements text")

            text = read_source_text(path)

            self.assertIn("PDF requirements text", text)

    def test_rejects_pdf_without_extractable_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "scanned.pdf"
            write_blank_pdf(path)

            with self.assertRaisesRegex(ValueError, "OCR"):
                read_source_text(path)

    def test_reads_two_column_pdf_in_visual_reading_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "paper.pdf"
            write_layout_pdf(path, two_columns=True)

            result = extract_pdf_text(path)

            self.assertEqual(result.page_layouts, ("two-column",))
            self.assertLess(result.text.index("FULL WIDTH PAPER TITLE"), result.text.index("LEFT COLUMN"))
            self.assertLess(result.text.index("LEFT COLUMN"), result.text.index("RIGHT COLUMN"))
            self.assertLess(result.text.index("RIGHT COLUMN"), result.text.index("FULL WIDTH FOOTER"))

    def test_keeps_single_column_pdf_in_vertical_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "report.pdf"
            write_layout_pdf(path, two_columns=False)

            result = extract_pdf_text(path)

            self.assertEqual(result.page_layouts, ("single-column",))
            self.assertLess(result.text.index("FULL WIDTH PAPER TITLE"), result.text.index("SINGLE COLUMN"))
            self.assertLess(result.text.index("SINGLE COLUMN"), result.text.index("FULL WIDTH FOOTER"))

    def test_reads_docx_headings_paragraphs_and_tables(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "source.docx"
            document = Document()
            document.add_heading("项目实施方案", level=1)
            document.add_paragraph("本项目用于提升文档修改效率。")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "阶段"
            table.cell(0, 1).text = "产出"
            table.cell(1, 0).text = "第一阶段"
            table.cell(1, 1).text = "原型"
            document.save(path)

            text = read_source_text(path)

            self.assertIn("# 项目实施方案", text)
            self.assertIn("本项目用于提升文档修改效率。", text)
            self.assertIn("| 阶段 | 产出 |", text)
            self.assertIn("| 第一阶段 | 原型 |", text)

    def test_writes_docx_from_markdown_like_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "final.docx"

            write_final_docx(
                "\n".join(
                    [
                        "# 修改稿",
                        "",
                        "这是正文段落。",
                        "",
                        "| 阶段 | 产出 |",
                        "| 第一阶段 | 原型 |",
                    ]
                ),
                output,
            )

            document = Document(output)
            self.assertEqual(document.paragraphs[0].text, "修改稿")
            self.assertTrue(document.paragraphs[0].style.name.startswith("Heading"))
            self.assertIn("这是正文段落。", [item.text for item in document.paragraphs])
            self.assertEqual(document.tables[0].cell(0, 0).text, "阶段")
            self.assertEqual(document.tables[0].cell(1, 1).text, "原型")


    def test_writes_markdown_formatting_as_real_docx_formatting(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "final.docx"

            write_final_docx(
                "\n".join(
                    [
                        "# Plan",
                        "",
                        "This project lasts **six months**.",
                        "",
                        "| Phase | Work |",
                        "| :--- | :--- |",
                        "| **Phase 1** | Line one<br>Line two |",
                    ]
                ),
                output,
            )

            document = Document(output)
            paragraph = next(item for item in document.paragraphs if "six months" in item.text)
            self.assertEqual(paragraph.text, "This project lasts six months.")
            self.assertTrue(any(run.text == "six months" and run.bold for run in paragraph.runs))
            table = document.tables[0]
            self.assertEqual(len(table.rows), 2)
            self.assertEqual(table.cell(1, 0).text, "Phase 1")
            self.assertNotIn("**", table.cell(1, 0).text)
            self.assertNotIn(":---", table.cell(1, 0).text)
            self.assertIn("Line one\nLine two", table.cell(1, 1).text)

    def test_skips_standalone_markdown_horizontal_rules(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "final.docx"

            write_final_docx(
                "\n".join(
                    [
                        "# Plan",
                        "---",
                        "First paragraph.",
                        "------",
                        "Second paragraph.",
                    ]
                ),
                output,
            )

            document = Document(output)
            paragraph_text = [item.text for item in document.paragraphs]
            self.assertNotIn("---", paragraph_text)
            self.assertNotIn("------", paragraph_text)
            self.assertIn("First paragraph.", paragraph_text)
            self.assertIn("Second paragraph.", paragraph_text)


if __name__ == "__main__":
    unittest.main()
