import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from office_revision.cli import main
from office_revision.config import ModelSettings
from office_revision.connection_test import ConnectionCheckResult


class CliTests(unittest.TestCase):
    def test_dry_run_writes_expected_output_files(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.md"
            requirements = root / "requirements.md"
            output = root / "output"
            source.write_text("原始项目方案。", encoding="utf-8")
            requirements.write_text("补充目标和进度安排。", encoding="utf-8")

            exit_code = main(
                [
                    "--source",
                    str(source),
                    "--requirements",
                    str(requirements),
                    "--output-dir",
                    str(output),
                    "--cycles",
                    "2",
                    "--dry-run",
                ]
            )

            self.assertEqual(exit_code, 0)
            final_text = (output / "final.md").read_text(encoding="utf-8")
            review_text = (output / "review.md").read_text(encoding="utf-8")
            run_log = json.loads((output / "run_log.json").read_text(encoding="utf-8"))

            self.assertIn("第 2 轮修改稿", final_text)
            self.assertIn("一、总体结论", review_text)
            self.assertIn("是否继续修改：是", review_text)
            self.assertEqual(run_log["cycles"], 2)
            self.assertEqual(run_log["actual_cycles"], 2)
            self.assertFalse(run_log["stopped_early"])
            self.assertEqual(len(run_log["passes"]), 2)
            self.assertIn("writer_instructions", run_log["passes"][0])

    def test_dry_run_with_docx_source_writes_final_docx(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.docx"
            requirements = root / "requirements.md"
            output = root / "output"
            document = Document()
            document.add_heading("项目实施方案", level=1)
            document.add_paragraph("原始项目方案。")
            document.save(source)
            requirements.write_text("补充目标和进度安排。", encoding="utf-8")

            exit_code = main(
                [
                    "--source",
                    str(source),
                    "--requirements",
                    str(requirements),
                    "--output-dir",
                    str(output),
                    "--cycles",
                    "1",
                    "--dry-run",
                ]
            )

            self.assertEqual(exit_code, 0)
            self.assertTrue((output / "final.md").exists())
            self.assertTrue((output / "final.docx").exists())
            final_doc = Document(output / "final.docx")
            self.assertIn("第 1 轮修改稿", [paragraph.text for paragraph in final_doc.paragraphs])

    def test_check_connections_does_not_require_source_or_requirements(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            config = Path(temp_dir) / "settings.env"
            config.write_text(
                "\n".join(
                    [
                        "WRITER_API_KEY=writer-key",
                        "WRITER_MODEL=writer-model",
                        "REVIEWER_API_KEY=reviewer-key",
                        "REVIEWER_MODEL=reviewer-model",
                    ]
                ),
                encoding="utf-8",
            )

            with patch(
                "office_revision.connection_test.check_all_connections",
                return_value=[
                    ConnectionCheckResult("WRITER", "writer-model", True, "ok"),
                    ConnectionCheckResult("REVIEWER", "reviewer-model", True, "ok"),
                ],
            ):
                exit_code = main(["--config", str(config), "--check-connections"])

            self.assertEqual(exit_code, 0)


    def test_uses_separate_default_output_directories_for_dry_run_and_real_runs(self):
        dry_run_args = main.__globals__["build_parser"]().parse_args(["--dry-run"])
        real_args = main.__globals__["build_parser"]().parse_args([])

        self.assertEqual(main.__globals__["default_output_dir"](dry_run_args), Path("outputs/demo/latest"))
        self.assertEqual(main.__globals__["default_output_dir"](real_args), Path("outputs/autogen/latest"))

    def test_default_run_output_dirs_include_timestamp_and_latest(self):
        dry_run_args = main.__globals__["build_parser"]().parse_args(["--dry-run"])
        real_args = main.__globals__["build_parser"]().parse_args([])
        explicit_args = main.__globals__["build_parser"]().parse_args(["--output-dir", "custom/out"])
        project_dir = Path("projects/Project_20260611")

        self.assertEqual(
            main.__globals__["default_run_output_dirs"](dry_run_args, "093000", project_dir=project_dir),
            [
                Path("projects/Project_20260611/dry_run_outputs/093000-pending-v1"),
                Path("projects/Project_20260611/dry_run_outputs/latest"),
            ],
        )
        self.assertEqual(
            main.__globals__["default_run_output_dirs"](real_args, "093000", project_dir=project_dir),
            [
                Path("projects/Project_20260611/outputs/093000-pending-v1"),
                Path("projects/Project_20260611/outputs/latest"),
            ],
        )
        self.assertEqual(
            main.__globals__["default_run_output_dirs"](explicit_args, "20260610_093000"),
            [Path("custom/out")],
        )

        self.assertEqual(
            main.__globals__["default_run_output_dirs"](
                dry_run_args,
                "093000",
                project_dir=project_dir,
                version=2,
            ),
            [
                Path("projects/Project_20260611/dry_run_outputs/093000-pending-v2"),
                Path("projects/Project_20260611/dry_run_outputs/latest"),
            ],
        )

    def test_locked_latest_directory_is_skipped_without_failure(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        root = Path(temp_dir.name)
        latest = root / "projects" / "Project_20260611" / "outputs" / "latest"
        timestamped = root / "projects" / "Project_20260611" / "outputs" / "093000-pending-v1"

        latest.mkdir(parents=True, exist_ok=True)
        with patch("office_revision.cli.shutil.rmtree", side_effect=PermissionError("locked")):
            self.assertFalse(main.__globals__["prepare_output_dir"](latest))
            self.assertTrue(main.__globals__["prepare_output_dir"](timestamped))

    def test_default_dry_run_creates_project_directory_outputs_and_input_snapshot(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            inputs = root / "inputs"
            projects = root / "projects"
            inputs.mkdir()
            (inputs / "source.md").write_text("Markdown source text.", encoding="utf-8")
            (inputs / "requirements.md").write_text("Improve it.", encoding="utf-8")

            with patch("office_revision.cli.DEFAULT_INPUT_DIR", inputs):
                exit_code = main(
                    [
                        "--projects-root",
                        str(projects),
                        "--project-title",
                        "Project Plan",
                        "--cycles",
                        "1",
                        "--dry-run",
                    ]
                )

            self.assertEqual(exit_code, 0)
            project_dirs = list(projects.glob("Project_Plan_*"))
            self.assertEqual(len(project_dirs), 1)
            project_dir = project_dirs[0]
            self.assertTrue((project_dir / "project.json").exists())
            self.assertTrue((project_dir / "inputs" / "source.md").exists())
            self.assertTrue((project_dir / "inputs" / "requirements.md").exists())
            self.assertTrue((project_dir / "dry_run_outputs" / "latest" / "final.md").exists())
            self.assertTrue((project_dir / "dry_run_outputs" / "latest" / "session_status.json").exists())

    def test_default_dry_run_increments_version_in_existing_project_directory(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            inputs = root / "inputs"
            projects = root / "projects"
            existing = projects / "Project_Plan_20260612" / "dry_run_outputs" / "093000-pending-v1"
            inputs.mkdir()
            existing.mkdir(parents=True)
            (inputs / "source.md").write_text("Markdown source text.", encoding="utf-8")
            (inputs / "requirements.md").write_text("Improve it.", encoding="utf-8")

            with patch("office_revision.cli.DEFAULT_INPUT_DIR", inputs), patch(
                "office_revision.cli.datetime"
            ) as fake_datetime:
                fake_datetime.now.return_value.strftime.side_effect = lambda fmt: {
                    "%Y%m%d": "20260612",
                    "%H%M%S": "094500",
                }[fmt]
                exit_code = main(
                    [
                        "--projects-root",
                        str(projects),
                        "--project-title",
                        "Project Plan",
                        "--cycles",
                        "1",
                        "--dry-run",
                    ]
                )

            self.assertEqual(exit_code, 0)
            project_dir = projects / "Project_Plan_20260612"
            self.assertTrue((project_dir / "dry_run_outputs" / "094500-pending-v2").exists())

    def test_choose_project_title_uses_llm_for_real_runs(self):
        args = main.__globals__["build_parser"]().parse_args(["--project-title-language", "zh"])
        reviewer_settings = ModelSettings(
            role="REVIEWER",
            api_key="key",
            base_url="",
            model="model",
        )

        with patch("office_revision.cli.generate_llm_project_title", return_value="项目实施方案修订"):
            title = main.__globals__["choose_project_title"](
                args,
                source_path=Path("inputs/source.docx"),
                source_text="项目实施方案正文",
                requirements="请修改",
                meeting_notes="",
                reviewer_settings=reviewer_settings,
            )

        self.assertEqual(title, "项目实施方案修订")

    def test_defaults_to_inputs_directory_for_daily_use(self):
        args = main.__globals__["build_parser"]().parse_args([])

        self.assertIsNone(args.source)
        self.assertIsNone(args.requirements)
        self.assertEqual(args.summary_mode, "rule")

    def test_accepts_llm_summary_mode(self):
        args = main.__globals__["build_parser"]().parse_args(["--summary-mode", "llm"])

        self.assertEqual(args.summary_mode, "llm")

    def test_docx_source_writes_round_drafts_and_reviews(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.docx"
            requirements = root / "requirements.md"
            output = root / "output"
            document = Document()
            document.add_heading("Project Plan", level=1)
            document.add_paragraph("Original draft.")
            document.save(source)
            requirements.write_text("Improve the plan.", encoding="utf-8")

            exit_code = main(
                [
                    "--source",
                    str(source),
                    "--requirements",
                    str(requirements),
                    "--output-dir",
                    str(output),
                    "--cycles",
                    "2",
                    "--dry-run",
                ]
            )

            self.assertEqual(exit_code, 0)
            self.assertTrue((output / "drafts" / "round_01_draft.md").exists())
            self.assertTrue((output / "drafts" / "round_01_draft.docx").exists())
            self.assertTrue((output / "drafts" / "round_02_draft.md").exists())
            self.assertTrue((output / "drafts" / "round_02_draft.docx").exists())
            self.assertTrue((output / "reviews" / "round_01_review.md").exists())
            self.assertTrue((output / "reviews" / "round_01_review.docx").exists())
            self.assertTrue((output / "reviews" / "round_02_review.md").exists())
            self.assertTrue((output / "reviews" / "round_02_review.docx").exists())
            draft_doc = Document(output / "drafts" / "round_01_draft.docx")
            review_doc = Document(output / "reviews" / "round_01_review.docx")
            self.assertTrue(any(paragraph.text for paragraph in draft_doc.paragraphs))
            self.assertTrue(any(paragraph.text for paragraph in review_doc.paragraphs))


    def test_runs_without_source_when_requirements_exist(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            inputs = root / "inputs"
            requirements = root / "requirements.md"
            output = root / "output"
            inputs.mkdir()
            requirements.write_text("Write a plan from scratch.", encoding="utf-8")

            with patch("office_revision.cli.DEFAULT_INPUT_DIR", inputs):
                exit_code = main(
                    [
                        "--requirements",
                        str(requirements),
                        "--output-dir",
                        str(output),
                        "--cycles",
                        "1",
                        "--dry-run",
                    ]
                )

            self.assertEqual(exit_code, 0)
            self.assertTrue((output / "final.md").exists())
            self.assertTrue((output / "final.docx").exists())
            self.assertTrue((output / "changes_summary.md").exists())
            self.assertTrue((output / "changes_summary.docx").exists())
            run_log = json.loads((output / "run_log.json").read_text(encoding="utf-8"))
            self.assertFalse(run_log["has_source"])
            self.assertIsNone(run_log["source_path"])

    def test_empty_source_and_empty_meeting_notes_are_treated_as_missing(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.md"
            requirements = root / "requirements.md"
            meeting_notes = root / "meeting_notes.md"
            output = root / "output"
            source.write_text("   \n", encoding="utf-8")
            requirements.write_text("Write a plan.", encoding="utf-8")
            meeting_notes.write_text("   \n", encoding="utf-8")

            exit_code = main(
                [
                    "--source",
                    str(source),
                    "--requirements",
                    str(requirements),
                    "--meeting-notes",
                    str(meeting_notes),
                    "--output-dir",
                    str(output),
                    "--cycles",
                    "1",
                    "--dry-run",
                ]
            )

            self.assertEqual(exit_code, 0)
            run_log = json.loads((output / "run_log.json").read_text(encoding="utf-8"))
            self.assertEqual(run_log["source_path"], str(source))
            self.assertEqual(run_log["meeting_notes_path"], str(meeting_notes))
            self.assertFalse(run_log["has_source"])
            self.assertFalse(run_log["has_meeting_notes"])

    def test_discovers_md_source_when_default_docx_is_missing(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            inputs = root / "inputs"
            output = root / "output"
            inputs.mkdir()
            (inputs / "source.md").write_text("Markdown source text.", encoding="utf-8")
            (inputs / "requirements.md").write_text("Improve it.", encoding="utf-8")

            with patch("office_revision.cli.DEFAULT_INPUT_DIR", inputs):
                exit_code = main(["--output-dir", str(output), "--cycles", "1", "--dry-run"])

            self.assertEqual(exit_code, 0)
            run_log = json.loads((output / "run_log.json").read_text(encoding="utf-8"))
            self.assertTrue(run_log["has_source"])
            self.assertEqual(run_log["source_path"], str(inputs / "source.md"))

    def test_missing_or_empty_requirements_stops_with_clear_error(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            empty_requirements = root / "requirements.md"
            empty_requirements.write_text(" \n", encoding="utf-8")

            with self.assertRaises(SystemExit) as raised:
                main(
                    [
                        "--requirements",
                        str(empty_requirements),
                        "--output-dir",
                        str(root / "output"),
                        "--dry-run",
                    ]
                )

            self.assertIn("requirements", str(raised.exception))

    def test_llm_summary_mode_falls_back_to_rule_summary_when_generation_fails(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            requirements = root / "requirements.md"
            output = root / "output"
            requirements.write_text("Write a plan.", encoding="utf-8")

            with patch(
                "office_revision.cli.generate_llm_changes_summary",
                side_effect=RuntimeError("summary model unavailable"),
            ):
                exit_code = main(
                    [
                        "--requirements",
                        str(requirements),
                        "--output-dir",
                        str(output),
                        "--cycles",
                        "1",
                        "--dry-run",
                        "--summary-mode",
                        "llm",
                    ]
                )

            self.assertEqual(exit_code, 0)
            self.assertTrue((output / "changes_summary.md").exists())
            run_log = json.loads((output / "run_log.json").read_text(encoding="utf-8"))
            self.assertEqual(run_log["summary_mode_requested"], "llm")
            self.assertEqual(run_log["summary_mode_used"], "rule")
            self.assertIn("summary model unavailable", run_log["summary_fallback_reason"])


if __name__ == "__main__":
    unittest.main()
